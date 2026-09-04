"""文档解析路由：按文件类型分派，统一输出文本 + 分页文本（用于按页分块/标注页码）。"""
from __future__ import annotations

import os
import re
import threading
import traceback
from dataclasses import dataclass, field
from typing import Any

from app.config import get_settings


@dataclass
class ParsedDocument:
    text: str
    page_count: int = 0
    pages: list[str] = field(default_factory=list)  # PDF 逐页；其他格式单元素
    metadata: dict[str, Any] = field(default_factory=dict)


def detect_kind(filename: str) -> str:
    ext = os.path.splitext(filename or "")[1].lower().lstrip(".")
    if ext == "pdf":
        return "pdf"
    if ext in ("docx", "doc"):
        return "docx"
    if ext in ("xlsx", "xls", "csv"):
        return "xlsx"
    if ext in ("md", "markdown"):
        return "md"
    if ext == "txt":
        return "txt"
    if ext in ("png", "jpg", "jpeg", "gif", "bmp"):
        return "image"
    return "unknown"


_docling_converter = None
_docling_lock = threading.Lock()



# ---------- 公式图片 → LaTeX（pix2text-mfr，TrOCR/optimum-onnx，走 HF/hf-mirror） ----------
_formula_ocr = None
_formula_ocr_lock = threading.Lock()


def _clean_formula_latex(latex: str) -> str:
    """规整 pix2text 输出的 LaTeX：压缩连续对齐间距宏（\qquad/\quad/\hspace）为单个空格，去多余空白。

    pix2text 常给公式尾部填一堆 \qquad 对齐（如 `... \qquad \qquad ... ( 2. 1 )`），会撑大 chunk 且占 token，
    这里只压缩间距、不动结构（保留 \\, \\begin, \\left/\right 等）。
    """
    s = latex or ""
    s = re.sub(r"(?:\\qquad|\\quad|\\hspace\*?\{[^}]*\})[ \t]*", " ", s)  # 压缩对齐间距宏
    s = re.sub(r"\\textcircled\s*\{\s*=\s*\}\s*", "@", s)                    # mAP@0.5 的 \textcircled{=}
    s = re.sub(r"\s*\{\s*", "{", s)                                            # { 前后去空格
    s = re.sub(r"\s*\}\s*", "}", s)                                            # } 前后去空格
    s = re.sub(r"[ \t]{2,}", " ", s)
    return s.strip()


def _get_formula_ocr():
    """惰性加载公式识别模型（首次会从 HF 下载一次，后常驻内存）。未安装/加载失败由调用方兜底。

    用 breezedeus/pix2text-mfr（TrOCR 架构，optimum-onnx 版）：pix2tex 模型从 GitHub 下载国内连不上，
    texify(VisionEncoderDecoder) 在 transformers 4.5x 报 .to_dict 错；这个 TrOCR 模型走 HF(hf-mirror) 且兼容当前 transformers。
    """
    global _formula_ocr
    if _formula_ocr is None:
        with _formula_ocr_lock:
            if _formula_ocr is None:
                from transformers import TrOCRProcessor
                from optimum.onnxruntime import ORTModelForVision2Seq

                processor = TrOCRProcessor.from_pretrained("breezedeus/pix2text-mfr")
                model = ORTModelForVision2Seq.from_pretrained("breezedeus/pix2text-mfr", use_cache=False)

                def ocr(img):  # 保持 ocr(img)->latex 接口不变
                    pixel_values = processor(images=img, return_tensors="pt").pixel_values
                    gen = model.generate(pixel_values)
                    return processor.tokenizer.batch_decode(gen, skip_special_tokens=True)[0]

                _formula_ocr = ocr
    return _formula_ocr


def _crop_formula_image(res, item, pdf, padding: int = 4, pad_left: int = 30):
    """裁剪公式区域为 PIL 图。优先 docling 自带 get_image（需 generate_page_images），否则 PyMuPDF 按 bbox 裁。

    pad_left 单独加大：docling 的公式 bbox 常漏掉公式最左侧的标识符（如 `L_{IoU}`/`L_{CE}`、`F^\prime`），
    而显示公式是居中排版、左侧是空白，多往左截些不会引入正文文字，能把标识符框进图里给 pix2text 识别。
    """
    from PIL import Image

    try:
        img = item.get_image(res.document)  # docling 内部已处理坐标系；需 page.image 非空
        if img is not None:
            return img
    except Exception:
        pass

    provs = getattr(item, "prov", None) or []
    if not provs:
        return None
    prov = provs[0]
    pgno = getattr(prov, "page_no", None)
    page = res.document.pages.get(pgno) if pgno is not None else None
    if page is None or page.size is None:
        return None
    box = prov.bbox.to_top_left_origin(page_height=page.size.height)
    try:
        pdf_page = pdf[pgno - 1]  # docling page_no 为 1-based
    except Exception:
        return None
    # docling 页尺寸与 pymupdf 若有细微差异，按比例缩放
    sx = pdf_page.rect.width / page.size.width if page.size.width else 1.0
    sy = pdf_page.rect.height / page.size.height if page.size.height else 1.0
    l = max(0.0, box.l * sx - pad_left)
    t = max(0.0, box.t * sy - padding)
    r = min(pdf_page.rect.width, box.r * sx + padding)
    b = min(pdf_page.rect.height, box.b * sy + padding)
    if r <= l or b <= t:
        return None
    import fitz
    pix = pdf_page.get_pixmap(dpi=300, clip=fitz.Rect(l, t, r, b), colorspace=fitz.csRGB, alpha=False)
    return Image.frombytes("RGB", (pix.width, pix.height), pix.samples)


def _enrich_formulas_with_ocr(res, pdf_path):
    """对 docling 标为公式但 text 为空的（not-decoded）公式，用公式识别模型(pix2text-mfr) 裁图识别为 LaTeX 注入 item.text。

    仅当 docling_formula_ocr 开启；公式识别模型缺装、加载失败或单条识别失败一律跳过（保留占位符），不影响主流程。
    带诊断日志（print 会被 verify_formula.py 捕获进报告），便于真机定位是"没下模型"还是"裁图坐标错"。
    """
    if not get_settings().docling_formula_ocr:
        return res
    try:
        from docling_core.types.doc.labels import DocItemLabel
    except Exception as e:
        print(f"[formula] docling label import 失败: {type(e).__name__}: {e}")
        return res
    items = []
    try:
        for item, _lvl in res.document.iterate_items():
            if item.label == DocItemLabel.FORMULA and not getattr(item, "text", ""):
                items.append(item)
    except Exception as e:
        print(f"[formula] iterate_items 遍历失败: {type(e).__name__}: {e}")
        return res
    if not items:
        print("[formula] 未找到 not-decoded 公式 item（无待识别）")
        return res
    print(f"[formula] 找到 {len(items)} 个 not-decoded 公式 item")
    try:
        ocr = _get_formula_ocr()
    except Exception as e:
        print(f"[formula] 公式识别模型加载失败: {type(e).__name__}: {e}")
        print("[formula] 加载异常回溯:\n" + traceback.format_exc())
        return res  # 识别模型缺装 / 模型下载/加载失败 → 跳过，保留占位符
    if ocr is None:
        return res
    print("[formula] 公式识别模型加载成功")
    import fitz
    try:
        pdf = fitz.open(os.path.abspath(pdf_path))
    except Exception as e:
        print(f"[formula] fitz 打开 PDF 失败: {type(e).__name__}: {e}")
        return res
    try:
        for i, item in enumerate(items):
            pgno = (item.prov[0].page_no if getattr(item, "prov", None) else "?")
            img = _crop_formula_image(res, item, pdf)
            if img is None:
                print(f"[formula]  #{i+1} p{pgno} 裁剪失败(bbox 不可用)")
                continue
            print(f"[formula]  #{i+1} p{pgno} 裁剪 OK 图像 {img.size[0]}x{img.size[1]}")
            try:
                latex = ocr(img)
            except Exception as e:
                print(f"[formula]  #{i+1} p{pgno} 识别异常: {type(e).__name__}: {e}")
                continue
            latex = _clean_formula_latex((latex or "").strip().strip("$").strip())  # 去定界符 + 规整间距
            if latex:
                item.text = latex
                print(f"[formula]  #{i+1} p{pgno} -> {latex[:50]}")
            else:
                print(f"[formula]  #{i+1} p{pgno} 识别为空")
    finally:
        pdf.close()
    return res


def _build_docling_converter():
    """构建 DocumentConverter（惰性 import；只构建一次，模型常驻内存）。

    - 后端：PyPdfium（绕开 docling_parse 的 additional.dat 缺失问题）。
    - OCR：关闭（文本型 PDF 不需要，OCR 极慢）。
    - device：优先 cuda（可用时），否则 cpu；可被 DOCLING_DEVICE 环境变量覆盖。
    """
    from docling.document_converter import DocumentConverter, PdfFormatOption  # 惰性
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions, TableStructureOptions
    from docling.datamodel.accelerator_options import AcceleratorOptions
    from docling.backend.pypdfium2_backend import PyPdfiumDocumentBackend

    dev = os.environ.get("DOCLING_DEVICE", "")
    if not dev:
        import torch
        dev = "cuda" if torch.cuda.is_available() else "cpu"
    pipe_opts = PdfPipelineOptions(
        do_ocr=False,  # 文本型 PDF 不需要 OCR（大提速）
        do_formula_enrichment=get_settings().docling_formula_enrichment,
        images_scale=get_settings().docling_images_scale,
        table_structure_options=TableStructureOptions(mode=get_settings().docling_table_mode),
        accelerator_options=AcceleratorOptions(device=dev, num_threads=8),
    )
    print(f"[doc] docling converter: device={dev} ocr=False")
    return DocumentConverter(
        format_options={
            InputFormat.PDF: PdfFormatOption(
                backend=PyPdfiumDocumentBackend,
                pipeline_options=pipe_opts,
            ),
        }
    )


def _get_docling_converter():
    global _docling_converter
    if _docling_converter is None:
        with _docling_lock:
            if _docling_converter is None:
                _docling_converter = _build_docling_converter()
    return _docling_converter


def _parse_pdf_docling(path: str):
    """Docling 解析 PDF（表格/版面/层级更强）。惰性 import，失败由调用方回退 PyMuPDF。

    用 PyPdfium 后端而非默认 docling_parse：docling_parse 的 Windows wheel 常缺
    pdf_resources/glyphs/standard/additional.dat（字体表）导致 RuntimeError；
    pdfium 不依赖该文件，版面/表格仍走 StandardPdfPipeline（需 onnxruntime）。
    复用模块级 converter：模型只加载一次，后续解析只需推理。

    按页导出 markdown（export_to_markdown(page_no=...)），使 chunk 带真实页码，
    引用溯源不再全显示"第1页"。
    """
    import os

    converter = _get_docling_converter()
    with _docling_lock:
        res = converter.convert(os.path.abspath(path))  # 绝对路径更稳

    _enrich_formulas_with_ocr(res, path)  # 公式图片→LaTeX（pix2text-mfr，可选，失败回退占位符）

    pages_map = getattr(res.document, "pages", None) or {}
    page_objs = list(pages_map.values()) if isinstance(pages_map, dict) else list(pages_map)
    page_count = len(page_objs) or 1
    pages: list[str] = []
    for p in page_objs:
        pno = getattr(p, "page_no", None)
        try:
            md_page = res.document.export_to_markdown(page_no=pno) if pno is not None else ""
        except Exception:  # noqa
            md_page = ""
        pages.append(md_page or "")  # 空页保留占位，保证 page_num 与真实页一致
    if not any(pg.strip() for pg in pages):  # 极端兜底：整篇单段
        pages = [res.document.export_to_markdown()]
    text = "\n\n".join(pages)
    return ParsedDocument(text=text, page_count=page_count, pages=pages,
                          metadata={"kind": "pdf", "parser": "docling"})


def _parse_pdf(path: str):
    import fitz  # PyMuPDF

    page_texts: list[str] = []
    with fitz.open(path) as doc:
        for page in doc:
            page_texts.append(page.get_text("text"))
    return "\n\n".join(page_texts), len(page_texts), page_texts


def _parse_docx(path: str):
    import docx

    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    text = "\n".join(parts)
    return text, len(d.paragraphs)


def _parse_xlsx(path: str):
    import openpyxl

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"# 工作表: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            vals = [str(c) for c in row if c is not None]
            if vals:
                parts.append(" | ".join(vals))
    return "\n".join(parts), len(wb.sheetnames)


def _parse_text(path: str):
    raw = open(path, encoding="utf-8", errors="ignore").read()
    return raw, raw.count("\n")


class ParserRouter:
    def parse(self, path: str, filename: str | None = None) -> ParsedDocument:
        name = filename or os.path.basename(path)
        kind = detect_kind(name)
        try:
            if kind == "pdf":
                if get_settings().parser_use_docling:
                    try:
                        return _parse_pdf_docling(path)
                    except Exception as e:  # noqa
                        import traceback
                        try:
                            import docling
                            _v = getattr(docling, "__version__", "?")
                        except Exception:
                            _v = "?"
                        print(f"[doc] docling({_v}) 失败，回退 PyMuPDF: {e}\n{traceback.format_exc()}")
                text, count, pages = _parse_pdf(path)
                return ParsedDocument(text=text, page_count=count, pages=pages, metadata={"kind": "pdf"})
            if kind == "docx":
                text, count = _parse_docx(path)
                return ParsedDocument(text=text, page_count=count, pages=[text], metadata={"kind": "docx"})
            if kind == "xlsx":
                text, count = _parse_xlsx(path)
                return ParsedDocument(text=text, page_count=count, pages=[text], metadata={"kind": "xlsx"})
            if kind in ("md", "txt"):
                text, count = _parse_text(path)
                return ParsedDocument(text=text, page_count=count, pages=[text], metadata={"kind": kind})
            if kind == "image":
                return ParsedDocument(text="", page_count=1, pages=[""], metadata={"kind": "image", "note": "需 OCR/多模态"})
            return ParsedDocument(text="", page_count=0, pages=[], metadata={"kind": "unknown", "error": "不支持的文件类型"})
        except ImportError as e:
            return ParsedDocument(text="", page_count=0, pages=[],
                                  metadata={"kind": kind, "error": f"缺少解析依赖: {e.name}"})
